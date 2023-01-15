from tkinter import *
from docx.shared import Inches
from listofmocktails import Mocktails
import docx
import os

root= Tk()

root.title("Select your ingredients")
root.iconbitmap("mocktail.ico")
root.geometry("350x750")

# created the checkboxes
var1 = IntVar()
Checkbutton(root, text = "Mint Leaves", variable = var1).grid(row=0,sticky=W)
var2 = IntVar()
Checkbutton(root, text = "Lime/Lemon Juice", variable = var2).grid(row=1,sticky=W)
var3 = IntVar()
Checkbutton(root, text = "Simple syrup or Sugar", variable = var3).grid(row=2,sticky=W)
var4 = IntVar()
Checkbutton(root, text = "Soda or Tonic/Sparkling water", variable = var4).grid(row=3,sticky=W)
var5 = IntVar()
Checkbutton(root, text = "Grenadine", variable = var5).grid(row=4,sticky=W)
var6 = IntVar()
Checkbutton(root, text = "Lemon-lime Soda (7up or Sprite)", variable = var6).grid(row=5,sticky=W)
var7 = IntVar()
Checkbutton(root, text = "Pineapple Juice", variable = var7).grid(row=6,sticky=W)
var8 = IntVar()
Checkbutton(root, text = "Orange Juice", variable = var8).grid(row=7,sticky=W)
var9 = IntVar()
Checkbutton(root, text = "Almond Syrup", variable = var9).grid(row=8,sticky=W)
var10 = IntVar()
Checkbutton(root, text = "Tomato Juice", variable = var10).grid(row=9,sticky=W)
var11 = IntVar()
Checkbutton(root, text = "Hot Pepper Sauce", variable = var11).grid(row=10,sticky=W)
var12 = IntVar()
Checkbutton(root, text = "Black Pepper", variable = var12).grid(row=11,sticky=W)
var13 = IntVar()
Checkbutton(root, text = "Horseradish", variable = var13).grid(row=12,sticky=W)
var14 = IntVar()
Checkbutton(root, text = "Peach Juice", variable = var14).grid(row=13,sticky=W)
var15 = IntVar()
Checkbutton(root, text = "Blueberries", variable = var15).grid(row=14,sticky=W)
var16 = IntVar()
Checkbutton(root, text = "Honey", variable = var16).grid(row=15,sticky=W)
var17 = IntVar()
Checkbutton(root, text = "Ginger Ale", variable = var17).grid(row=16,sticky=W)
var18 = IntVar()
Checkbutton(root, text = "Grapefruit Juice", variable = var18).grid(row=17,sticky=W)
var19 = IntVar()
Checkbutton(root, text = "Ginger Beer", variable = var19).grid(row=18,sticky=W)
var20 = IntVar()
Checkbutton(root, text = "Cucumber", variable = var20).grid(row=19,sticky=W)
var21 = IntVar()
Checkbutton(root, text = "Strawberries", variable = var21).grid(row=20,sticky=W)
var22 = IntVar()
Checkbutton(root, text = "Jalapeno", variable = var22).grid(row=21,sticky=W)
var23 = IntVar()
Checkbutton(root, text = "Peaches", variable = var23).grid(row=22,sticky=W)
var24 = IntVar()
Checkbutton(root, text = "Apple Juice", variable = var24).grid(row=23,sticky=W)
var25 = IntVar()
Checkbutton(root, text = "Cranberry Juice", variable = var25).grid(row=24,sticky=W)
var26 = IntVar()
Checkbutton(root, text = "Pomegranate", variable = var26).grid(row=25,sticky=W)
Button(root,text = "OK", command=root.destroy, width=10).grid(row=26, sticky=W)

root.mainloop()

youringredients= []
listofingredients = [[var1, "mint leaves"],[var2,"lime juice"],[var3,"simple syrup"],[var4,"soda"],[var5, "grenadine"],[var6,"sprite"],[var7,"pineapple juice"],[var8,"orange juice"],[var9,"almond syrup"],[var10,"tomato juice"],
                     [var11,"hot pepper sauce"],[var12,"black pepper"],[var13,"horseradish"],[var14,"peach juice"],[var15,"blueberries"],[var16,"honey"],[var17,"ginger ale"],[var18,"grapefruit juice"],[var19,"ginger beer"],
                     [var20,"cucumber"],[var21,"strawberries"],[var22,"jalapeno"],[var23,"peaches"],[var24,"apple juice"],[var25,"cranberry juice"],[var26,"pomegranate"]]

for ingredient in range(len(listofingredients)):
    if listofingredients[ingredient][0].get()==1:
        youringredients.append(listofingredients[ingredient][1])

print(f"Your Ingredients are: ", end=" ")
print(*youringredients, sep=", ")

# new list of mocktails you can make (missing less then or = 2)
mktails_you_can_make = []

# makeadrink adds all the drinks you can make to the list mktails_you_can_make
def makeadrink(mocktail):
    missingingredients = list(set(mocktail[1]) - set(youringredients))
    if len(missingingredients) <= 2:
        mktails_you_can_make.append([mocktail[0], mocktail[2][0], missingingredients]) #adds the mocktail name, instructions and missing ingredients respectively.
    return mktails_you_can_make

# Sorting sorts a list according to length of 2nd item in the list
def Sorting(sub_li):
    sub_li.sort(key=lambda x: len(x[2]))
    return sub_li

def main():
    for mocktail in Mocktails:
        makeadrink(mocktail)
    Sorting(mktails_you_can_make)

    for mocktail in mktails_you_can_make:
        mocktail[2] = ', '.join(mocktail[2])
        if(len(mocktail[2]) == 0):
            mocktail[2] = 'No missing ingredients!'
    print(mktails_you_can_make)

    doc = docx.Document()
    doc.add_heading('List of Mocktails to Make', 0)

    for mocktail in mktails_you_can_make:
        doc.add_heading(mocktail[0])
        doc.add_picture(f'{mocktail[0]}.jpg', width=Inches(1), height=Inches(1))
        doc.add_paragraph(f"Missing Ingredients: {mocktail[2]}")
        doc.add_paragraph(f'Instructions: {mocktail[1]}')

    doc.save('mocktails.docx')

main()
os.system('mocktails.docx')

